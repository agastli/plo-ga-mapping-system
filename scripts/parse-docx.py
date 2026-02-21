#!/usr/bin/env python3
"""
Enhanced Word Document Parser for PLO-GA Mapping
Handles multiple document formats
"""

import sys
import json
import re
from docx import Document

def extract_program_info(doc):
    """Extract program information from the document"""
    program_info = {
        "programNameEn": "",
        "programNameAr": "",
        "departmentEn": "",
        "departmentAr": "",
        "collegeEn": "",
        "collegeAr": "",
        "language": "en"
    }
    
    for para in doc.paragraphs:
        text = para.text.strip()
        
        # Extract program name
        if "Program:" in text or "Program Name:" in text or "البرنامج:" in text:
            program_info["programNameEn"] = text.split("Program")[-1].split(":")[-1].strip()
            program_info["programNameAr"] = text.split("البرنامج:")[-1].strip() if "البرنامج:" in text else ""
        
        # Extract department
        if "Department:" in text or "القسم:" in text:
            program_info["departmentEn"] = text.split("Department:")[-1].strip() if "Department:" in text else ""
            program_info["departmentAr"] = text.split("القسم:")[-1].strip() if "القسم:" in text else ""
        
        # Extract college
        if "College:" in text or "الكلية:" in text:
            program_info["collegeEn"] = text.split("College:")[-1].strip() if "College:" in text else ""
            program_info["collegeAr"] = text.split("الكلية:")[-1].strip() if "الكلية:" in text else ""
    
    # Detect language
    if program_info["programNameAr"] or program_info["departmentAr"]:
        program_info["language"] = "ar"
    
    return program_info

def extract_plos(doc):
    """Extract Program Learning Outcomes - handles multiple formats"""
    plos = []
    in_plo_section = False
    plo_counter = 1
    
    # Patterns for different PLO formats
    plo_with_code = re.compile(r'^(PLO\s*\d+)[:\s]*(.+)', re.IGNORECASE)
    numbered_list = re.compile(r'^(\d+)[.)]\s+(.+)')
    
    for para in doc.paragraphs:
        text = para.text.strip()
        
        # Start PLO section
        if "Program Learning Outcomes" in text or "مخرجات التعلم" in text:
            in_plo_section = True
            continue
        
        # Stop at mapping section
        if in_plo_section and ("Mapping" in text or "المواءمة" in text):
            break
        
        if in_plo_section and text:
            code = None
            description = None
            
            # Try format: "PLO1: description"
            match = plo_with_code.match(text)
            if match:
                code = match.group(1).replace(" ", "").upper()
                description = match.group(2).strip()
            else:
                # Try format: "1. description"
                match = numbered_list.match(text)
                if match:
                    number = match.group(1)
                    code = f"PLO{number}"
                    description = match.group(2).strip()
            
            if code and description:
                is_arabic = bool(re.search(r'[\u0600-\u06FF]', description))
                plos.append({
                    "code": code,
                    "descriptionEn": description if not is_arabic else "",
                    "descriptionAr": description if is_arabic else "",
                    "sortOrder": plo_counter
                })
                plo_counter += 1
    
    return plos

def extract_mapping_matrix(doc):
    """Extract mapping matrix from tables"""
    mappings = {}
    found_valid_table = False
    
    for table in doc.tables:
        # Check if this is the mapping matrix table
        header_text = " ".join([cell.text.strip() for cell in table.rows[0].cells]).lower()
        if "plo" not in header_text:
            continue
        
        found_valid_table = True
        
        # Extract PLO codes from header row
        header_cells = [cell.text.strip() for cell in table.rows[0].cells]
        plo_indices = {}
        for i, text in enumerate(header_cells):
            if re.match(r'PLO\d+', text, re.IGNORECASE):
                plo_indices[i] = text.upper()
        
        if not plo_indices:
            continue
        
        # Extract mappings from data rows
        for row in table.rows[1:]:
            cells = [cell.text.strip() for cell in row.cells]
            if len(cells) < 3:
                continue
            
            # Extract competency code (usually in second column)
            competency_match = None
            for cell_text in cells[:2]:
                competency_match = re.search(r'(C\d+-\d+)', cell_text)
                if competency_match:
                    break
            
            if not competency_match:
                continue
            
            competency_code = competency_match.group(1)
            
            # Extract weights for each PLO
            for col_idx, plo_code in plo_indices.items():
                if col_idx < len(cells):
                    weight_text = cells[col_idx].strip()
                    try:
                        weight = float(weight_text) if weight_text else 0.0
                        key = f"{plo_code}_{competency_code}"
                        mappings[key] = {
                            "ploCode": plo_code,
                            "competencyCode": competency_code,
                            "weight": weight
                        }
                    except ValueError:
                        pass
    
    # Validate that we found a proper mapping table
    if not found_valid_table:
        raise ValueError(
            "No valid mapping table found. The document must contain a table with:\n"
            "- Column 1: Graduate Attributes (GA1, GA2, etc.)\n"
            "- Column 2: Supporting Competencies (C1-1, C1-2, etc.)\n"
            "- Columns 3+: PLO1, PLO2, PLO3, etc.\n\n"
            "Please ensure your document follows the template structure with a proper table format."
        )
    
    if not mappings:
        raise ValueError(
            "No mapping data found in the table. Please ensure:\n"
            "- The table has competency codes (C1-1, C1-2, etc.) in the second column\n"
            "- Weight values (0.0 to 1.0) are entered in the PLO columns\n"
            "- The document follows the exact template structure"
        )
    
    return list(mappings.values())

def extract_justifications(doc):
    """Extract competency-based justifications - one justification per competency"""
    justifications = []
    in_justification_section = False
    current_ga = None
    current_competency = None
    justification_buffer = []
    
    # Competency patterns
    ga_pattern = re.compile(r'^(GA\s*\d+)', re.IGNORECASE)
    competency_pattern = re.compile(r'^(C\d+-\d+)[:\s]*(.+)', re.IGNORECASE)
    numbered_list = re.compile(r'^(\d+)[.)]\s+(.+)')
    
    for para in doc.paragraphs:
        text = para.text.strip()
        
        # Start justification section
        if "Justifications" in text or "مبررات" in text or "Justification" in text:
            in_justification_section = True
            continue
        
        if not in_justification_section or not text:
            continue
        
        # Check for GA header
        ga_match = ga_pattern.match(text)
        if ga_match:
            current_ga = ga_match.group(1).replace(" ", "").upper()
            continue
        
        # Check for competency header with code (e.g., "C1-1:" or "C1-1 -")
        comp_match = competency_pattern.match(text)
        if comp_match:
            # Save previous justification
            if current_ga and current_competency and justification_buffer:
                justifications.append({
                    "gaCode": current_ga,
                    "competencyCode": current_competency,
                    "textEn": " ".join(justification_buffer),
                    "textAr": ""
                })
            
            current_competency = comp_match.group(1).upper()
            # Check if justification starts on same line
            rest_of_line = comp_match.group(2).strip()
            if rest_of_line:
                justification_buffer = [rest_of_line]
            else:
                justification_buffer = []
            continue
        
        # Accumulate justification text
        if current_ga and current_competency and text:
            justification_buffer.append(text)
    
    # Save last justification
    if current_ga and current_competency and justification_buffer:
        justifications.append({
            "gaCode": current_ga,
            "competencyCode": current_competency,
            "textEn": " ".join(justification_buffer),
            "textAr": ""
        })
    
    return justifications

def parse_document(file_path):
    """Main parsing function"""
    try:
        doc = Document(file_path)
        
        result = {
            "success": True,
            "data": {
                "programInfo": extract_program_info(doc),
                "plos": extract_plos(doc),
                "mappings": extract_mapping_matrix(doc),
                "justifications": extract_justifications(doc)
            }
        }
        
        return result
    
    except Exception as e:
        return {
            "success": False,
            "error": str(e)
        }

if __name__ == "__main__":
    if len(sys.argv) < 2:
        print(json.dumps({"success": False, "error": "No file path provided"}))
        sys.exit(1)
    
    file_path = sys.argv[1]
    result = parse_document(file_path)
    print(json.dumps(result, ensure_ascii=False, indent=2))
