#!/usr/bin/env python3
"""
Enhanced Word Document Parser for PLO-GA Mapping
Handles multiple document formats
"""

import sys
import json
import re
from docx import Document

def extract_program_info(doc):
    """Extract program information from the document"""
    program_info = {
        "programNameEn": "",
        "programNameAr": "",
        "departmentEn": "",
        "departmentAr": "",
        "collegeEn": "",
        "collegeAr": "",
        "language": "en"
    }
    
    for para in doc.paragraphs:
        text = para.text.strip()
        
        # Extract program name
        if "Program:" in text or "Program Name:" in text or "البرنامج:" in text:
            program_info["programNameEn"] = text.split("Program")[-1].split(":")[-1].strip()
            program_info["programNameAr"] = text.split("البرنامج:")[-1].strip() if "البرنامج:" in text else ""
        
        # Extract department
        if "Department:" in text or "القسم:" in text:
            program_info["departmentEn"] = text.split("Department:")[-1].strip() if "Department:" in text else ""
            program_info["departmentAr"] = text.split("القسم:")[-1].strip() if "القسم:" in text else ""
        
        # Extract college
        if "College:" in text or "الكلية:" in text:
            program_info["collegeEn"] = text.split("College:")[-1].strip() if "College:" in text else ""
            program_info["collegeAr"] = text.split("الكلية:")[-1].strip() if "الكلية:" in text else ""
    
    # Detect language
    if program_info["programNameAr"] or program_info["departmentAr"]:
        program_info["language"] = "ar"
    
    return program_info

def extract_plos(doc):
    """Extract Program Learning Outcomes - handles multiple formats"""
    plos = []
    in_plo_section = False
    
    # Patterns for different PLO formats
    plo_with_code = re.compile(r'^(PLO\s*\d+)[:\s]*(.+)', re.IGNORECASE)
    numbered_list = re.compile(r'^(\d+)[.)]\s+(.+)')
    
    for para in doc.paragraphs:
        text = para.text.strip()
        
        # Start PLO section
        if "Program Learning Outcomes" in text or "مخرجات التعلم" in text:
            in_plo_section = True
            continue
        
        # Stop at mapping section
        if in_plo_section and ("Mapping" in text or "المواءمة" in text):
            break
        
        if in_plo_section and text:
            code = None
            description = None
            
            # Try format: "PLO1: description"
            match = plo_with_code.match(text)
            if match:
                code = match.group(1).replace(" ", "").upper()
                description = match.group(2).strip()
            else:
                # Try format: "1. description"
                match = numbered_list.match(text)
                if match:
                    number = match.group(1)
                    code = f"PLO{number}"
                    description = match.group(2).strip()
            
            if code and description:
                is_arabic = bool(re.search(r'[\u0600-\u06FF]', description))
                plos.append({
                    "code": code,
                    "descriptionEn": description if not is_arabic else "",
                    "descriptionAr": description if is_arabic else ""
                })
    
    return plos

def extract_mapping_matrix(doc):
    """Extract mapping matrix from tables"""
    mappings = {}
    
    for table in doc.tables:
        # Check if this is the mapping matrix table
        header_text = " ".join([cell.text.strip() for cell in table.rows[0].cells]).lower()
        if "plo" not in header_text and "graduate" not in header_text:
            continue
        
        # Extract PLO codes from header
        plo_codes = []
        for cell in table.rows[0].cells:
            text = cell.text.strip()
            if re.match(r'PLO\s*\d+', text, re.IGNORECASE):
                plo_codes.append(text.replace(" ", "").upper())
            elif re.match(r'^\d+$', text):
                plo_codes.append(f"PLO{text}")
        
        if not plo_codes:
            continue
        
        # Extract mappings
        for row in table.rows[1:]:
            cells = [cell.text.strip() for cell in row.cells]
            if len(cells) < 2:
                continue
            
            # Extract competency code from first column
            competency_match = re.search(r'(C\d+-\d+)', cells[0])
            if not competency_match:
                continue
            
            competency_code = competency_match.group(1)
            
            # Extract weights for each PLO
            for i, plo_code in enumerate(plo_codes):
                if i + 1 < len(cells):
                    weight_text = cells[i + 1].strip()
                    try:
                        weight = float(weight_text) if weight_text else 0.0
                        key = f"{plo_code}_{competency_code}"
                        mappings[key] = {
                            "ploCode": plo_code,
                            "competencyCode": competency_code,
                            "weight": weight
                        }
                    except ValueError:
                        pass
    
    return list(mappings.values())

def extract_justifications(doc):
    """Extract justifications - handles multiple formats"""
    justifications = []
    in_justification_section = False
    current_ga = None
    current_competency = None
    current_competency_name = None
    justification_buffer = []
    
    # GA patterns
    ga_patterns = [
        (r'GA1[:\s]*Competent', 'GA1'),
        (r'GA2[:\s]*Life-long Learner', 'GA2'),
        (r'GA3[:\s]*Well Rounded', 'GA3'),
        (r'GA4[:\s]*Ethically.*Responsible', 'GA4'),
        (r'GA5[:\s]*Entrepreneurial', 'GA5'),
        (r'Graduate Attribute 1[:\s]*Competent', 'GA1'),
        (r'Graduate Attribute 2[:\s]*Life-long Learner', 'GA2'),
        (r'Graduate Attribute 3[:\s]*Well Rounded', 'GA3'),
        (r'Graduate Attribute 4[:\s]*Ethically.*Responsible', 'GA4'),
        (r'Graduate Attribute 5[:\s]*Entrepreneurial', 'GA5'),
    ]
    
    # Competency patterns
    competency_with_justification = re.compile(r'^(C\d+-\d+)\s*[–-]\s*(.+?):\s*(.+)', re.DOTALL)
    competency_header_only = re.compile(r'^(C\d+-\d+)[:\s]+(.+?)$')
    
    for para in doc.paragraphs:
        text = para.text.strip()
        
        # Start justification section
        if "Justifications for Mapping" in text or "مبررات المواءمة" in text:
            in_justification_section = True
            continue
        
        if not in_justification_section or not text:
            continue
        
        # Check for GA header
        for pattern, ga_code in ga_patterns:
            if re.search(pattern, text, re.IGNORECASE):
                # Save previous justification if exists
                if current_competency and justification_buffer:
                    justifications.append({
                        "gaCode": current_ga,
                        "competencyCode": current_competency,
                        "competencyName": current_competency_name,
                        "textEn": " ".join(justification_buffer),
                        "textAr": ""
                    })
                    justification_buffer = []
                
                current_ga = ga_code
                current_competency = None
                break
        
        # Check for competency with justification on same line
        match = competency_with_justification.match(text)
        if match:
            # Save previous justification
            if current_competency and justification_buffer:
                justifications.append({
                    "gaCode": current_ga,
                    "competencyCode": current_competency,
                    "competencyName": current_competency_name,
                    "textEn": " ".join(justification_buffer),
                    "textAr": ""
                })
            
            current_competency = match.group(1)
            current_competency_name = match.group(2).strip()
            justification_buffer = [match.group(3).strip()]
            continue
        
        # Check for competency header only (justification on next lines)
        match = competency_header_only.match(text)
        if match:
            # Save previous justification
            if current_competency and justification_buffer:
                justifications.append({
                    "gaCode": current_ga,
                    "competencyCode": current_competency,
                    "competencyName": current_competency_name,
                    "textEn": " ".join(justification_buffer),
                    "textAr": ""
                })
            
            current_competency = match.group(1)
            current_competency_name = match.group(2).strip()
            justification_buffer = []
            continue
        
        # Accumulate justification text
        if current_competency and text:
            justification_buffer.append(text)
    
    # Save last justification
    if current_competency and justification_buffer:
        justifications.append({
            "gaCode": current_ga,
            "competencyCode": current_competency,
            "competencyName": current_competency_name,
            "textEn": " ".join(justification_buffer),
            "textAr": ""
        })
    
    return justifications

def parse_document(file_path):
    """Main parsing function"""
    try:
        doc = Document(file_path)
        
        result = {
            "success": True,
            "data": {
                "programInfo": extract_program_info(doc),
                "plos": extract_plos(doc),
                "mappings": extract_mapping_matrix(doc),
                "justifications": extract_justifications(doc)
            }
        }
        
        return result
    
    except Exception as e:
        return {
            "success": False,
            "error": str(e)
        }

if __name__ == "__main__":
    if len(sys.argv) < 2:
        print(json.dumps({"success": False, "error": "No file path provided"}))
        sys.exit(1)
    
    file_path = sys.argv[1]
    result = parse_document(file_path)
    print(json.dumps(result, ensure_ascii=False, indent=2))
