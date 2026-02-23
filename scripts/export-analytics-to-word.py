#!/usr/bin/env python3
"""
Export analytics data to Word format
"""

import sys
import json
import os
import base64
import io
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from PIL import Image as PILImage

def add_horizontal_line(doc):
    """Add a decorative horizontal line"""
    p = doc.add_paragraph()
    p_fmt = p.paragraph_format
    p_fmt.space_before = Pt(6)
    p_fmt.space_after = Pt(6)
    
    pPr = p._element.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '12')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), '8B1538')
    pBdr.append(bottom)
    pPr.append(pBdr)

def create_analytics_word(data, output_path, logo_path):
    """Generate Word analytics report"""
    
    doc = Document()
    
    # Set page to portrait A4
    section = doc.sections[0]
    section.page_height = Inches(11.69)
    section.page_width = Inches(8.27)
    section.left_margin = Inches(0.75)
    section.right_margin = Inches(0.75)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    
    # Add logo
    if logo_path and os.path.exists(logo_path):
        try:
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run()
            run.add_picture(logo_path, width=Inches(2.8))
        except Exception as e:
            print(f"Warning: Could not add logo: {e}", file=sys.stderr)
    
    # Office name
    p = doc.add_paragraph("Academic Planning & Quality Assurance Office")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.runs[0]
    run.font.size = Pt(12)
    run.font.color.rgb = RGBColor(102, 102, 102)
    
    # Decorative line
    add_horizontal_line(doc)
    
    # Title
    title = doc.add_heading(data.get('title', 'Analytics Report'), level=1)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.runs[0]
    run.font.color.rgb = RGBColor(139, 21, 56)
    run.font.size = Pt(24)
    
    # Timestamp
    if 'timestamp' in data:
        p = doc.add_paragraph(f"Generated on: {data['timestamp']}")
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.runs[0]
        run.font.size = Pt(10)
        run.font.color.rgb = RGBColor(102, 102, 102)
    
    doc.add_paragraph()  # Spacer
    
    # Key Metrics section
    if 'metrics' in data and len(data['metrics']) > 0:
        heading = doc.add_heading('Key Metrics', level=2)
        run = heading.runs[0]
        run.font.color.rgb = RGBColor(139, 21, 56)
        
        # Create metrics table
        table = doc.add_table(rows=len(data['metrics']) + 1, cols=2)
        table.style = 'Light Grid Accent 1'
        
        # Header row
        header_cells = table.rows[0].cells
        header_cells[0].text = "Metric"
        header_cells[1].text = "Value"
        
        for cell in header_cells:
            cell.paragraphs[0].runs[0].font.bold = True
            cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)
            shading_elm = OxmlElement('w:shd')
            shading_elm.set(qn('w:fill'), '8B1538')
            cell._element.get_or_add_tcPr().append(shading_elm)
        
        # Data rows
        for idx, metric in enumerate(data['metrics'], start=1):
            row_cells = table.rows[idx].cells
            row_cells[0].text = metric['label']
            row_cells[1].text = str(metric['value'])
        
        # Set column widths
        for row in table.rows:
            row.cells[0].width = Inches(4)
            row.cells[1].width = Inches(2)
        
        doc.add_paragraph()  # Spacer
    
    # Add chart images if provided (each as separate section)
    if 'chart_images' in data and len(data['chart_images']) > 0:
        doc.add_page_break()  # Start charts on new page
        for chart in data['chart_images']:
            if 'path' in chart and os.path.exists(chart['path']):
                # Add chart title
                heading = doc.add_heading(chart.get('title', 'Chart'), level=2)
                run = heading.runs[0]
                run.font.color.rgb = RGBColor(139, 21, 56)
                
                try:
                    # Add chart image
                    p = doc.add_paragraph()
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    run = p.add_run()
                    run.add_picture(chart['path'], width=Inches(6.5))
                    
                    doc.add_paragraph()  # Spacer
                except Exception as e:
                    print(f"Warning: Could not add chart image {chart.get('title', '')}: {e}", file=sys.stderr)
    
    # Add page break before data table
    doc.add_page_break()
    
    # Data table
    if 'table_data' in data and len(data['table_data']) > 0:
        heading = doc.add_heading('Detailed Data', level=2)
        run = heading.runs[0]
        run.font.color.rgb = RGBColor(139, 21, 56)
        
        table_data = data['table_data']
        table = doc.add_table(rows=len(table_data), cols=len(table_data[0]))
        table.style = 'Light Grid Accent 1'
        
        # Header row
        for col_idx, header in enumerate(table_data[0]):
            cell = table.rows[0].cells[col_idx]
            cell.text = header
            cell.paragraphs[0].runs[0].font.bold = True
            cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            shading_elm = OxmlElement('w:shd')
            shading_elm.set(qn('w:fill'), '8B1538')
            cell._element.get_or_add_tcPr().append(shading_elm)
        
        # Data rows
        for row_idx, row_data in enumerate(table_data[1:], start=1):
            for col_idx, value in enumerate(row_data):
                cell = table.rows[row_idx].cells[col_idx]
                cell.text = str(value)
                cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Add page numbers to footer
    section = doc.sections[0]
    footer = section.footer
    p = footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = p.add_run()
    fldChar1 = OxmlElement('w:fldChar')
    fldChar1.set(qn('w:fldCharType'), 'begin')
    run._element.append(fldChar1)
    
    instrText = OxmlElement('w:instrText')
    instrText.text = "PAGE"
    run._element.append(instrText)
    
    fldChar2 = OxmlElement('w:fldChar')
    fldChar2.set(qn('w:fldCharType'), 'end')
    run._element.append(fldChar2)
    
    # Save document
    doc.save(output_path)
    
    return {"success": True, "path": output_path}

if __name__ == "__main__":
    try:
        if len(sys.argv) < 2:
            print(json.dumps({"success": False, "error": "No input file provided"}))
            sys.exit(1)
        
        temp_file = sys.argv[1]
        with open(temp_file, 'r', encoding='utf-8') as f:
            input_data = json.load(f)
        
        data = input_data['data']
        output_path = input_data['output_path']
        logo_path = input_data.get('logo_path', '')
        
        result = create_analytics_word(data, output_path, logo_path)
        print(json.dumps(result))
        
    except Exception as e:
        print(json.dumps({"success": False, "error": str(e)}))
        sys.exit(1)
