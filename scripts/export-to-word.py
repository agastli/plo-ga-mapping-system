#!/usr/bin/env python3
"""
Export PLO-GA mapping data to Word document format
Generates a document matching the PDF format exactly
"""

import sys
import json
from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.enum.section import WD_ORIENT

# Qatar University Colors
QU_MAROON = RGBColor(139, 21, 56)  # #8B1538
QU_GOLD = RGBColor(212, 175, 55)   # #D4AF37
LIGHT_GOLD = RGBColor(200, 168, 130)  # #C8A882
LIGHT_GRAY = RGBColor(245, 245, 245)  # #F5F5F5
MEDIUM_GRAY = RGBColor(224, 224, 224)  # #E0E0E0
DARK_GRAY = RGBColor(85, 85, 85)  # #555555

def set_cell_background(cell, color):
    """Set cell background color"""
    shading_elm = OxmlElement('w:shd')
    shading_elm.set(qn('w:fill'), color)
    cell._element.get_or_add_tcPr().append(shading_elm)

def add_horizontal_line(doc, color_rgb, width_pt=2):
    """Add a horizontal line"""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    
    pPr = p._element.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), str(width_pt * 4))  # Size in eighths of a point
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), '%02x%02x%02x' % color_rgb)
    pBdr.append(bottom)
    pPr.append(pBdr)
    
    return p

def create_mapping_document(data):
    """Create Word document matching PDF format exactly"""
    doc = Document()
    
    # Set to landscape A4 with same margins as PDF
    section = doc.sections[0]
    section.orientation = WD_ORIENT.LANDSCAPE
    section.page_width = Inches(11.69)  # A4 landscape width
    section.page_height = Inches(8.27)  # A4 landscape height
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(0.75)
    section.right_margin = Inches(0.75)
    
    # Add QU logo at the top (centered)
    if data.get('logo_path'):
        try:
            import os
            logo_path = data['logo_path']
            print(f"DEBUG: Logo path: {logo_path}", file=sys.stderr)
            print(f"DEBUG: Logo exists: {os.path.exists(logo_path)}", file=sys.stderr)
            if os.path.exists(logo_path):
                # QU logo aspect ratio is 4.67:1 (2048x439 pixels)
                # Set width and let python-docx calculate height to preserve aspect ratio
                doc.add_picture(logo_path, width=Inches(2.5))
                last_paragraph = doc.paragraphs[-1]
                last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                last_paragraph.paragraph_format.space_after = Pt(12)
                print(f"DEBUG: Logo added successfully", file=sys.stderr)
            else:
                print(f"ERROR: Logo file not found at {logo_path}", file=sys.stderr)
        except Exception as e:
            print(f"ERROR: Could not add logo: {e}", file=sys.stderr)
            import traceback
            traceback.print_exc(file=sys.stderr)
    
    # Add "Academic Planning & Quality Assurance Office" under logo
    office_para = doc.add_paragraph('Academic Planning & Quality Assurance Office')
    office_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    office_para.runs[0].font.size = Pt(10)
    office_para.runs[0].font.italic = True
    office_para.runs[0].font.color.rgb = RGBColor(128, 128, 128)
    office_para.paragraph_format.space_after = Pt(16)
    
    # Add decorative lines
    add_horizontal_line(doc, (139, 21, 56), 2)  # Maroon line
    add_horizontal_line(doc, (212, 175, 55), 0.5)  # Gold line
    
    doc.add_paragraph().paragraph_format.space_after = Pt(20)
    
    # Add title
    title = doc.add_heading(data['program_name'], level=1)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.runs[0].font.color.rgb = QU_MAROON
    title.runs[0].font.size = Pt(24)
    title.paragraph_format.space_after = Pt(6)
    
    # Add subtitle
    subtitle = doc.add_paragraph('Program Learning Outcomes to Graduate Attributes Mapping')
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.runs[0].font.size = Pt(14)
    subtitle.runs[0].font.color.rgb = DARK_GRAY
    subtitle.paragraph_format.space_after = Pt(20)
    
    # Add decorative lines again
    add_horizontal_line(doc, (139, 21, 56), 2)
    add_horizontal_line(doc, (212, 175, 55), 0.5)
    
    doc.add_paragraph().paragraph_format.space_after = Pt(20)
    
    # Add program information in a table
    info_table = doc.add_table(rows=4, cols=2)
    info_table.style = 'Table Grid'
    
    info_data = [
        ('College:', data['college_name']),
        ('Department:', data['department_name']),
        ('Language:', data['language']),
        ('Last Updated:', data.get('last_updated', 'N/A'))
    ]
    
    for idx, (label, value) in enumerate(info_data):
        info_table.cell(idx, 0).text = label
        info_table.cell(idx, 1).text = value
        
        # Style label cells
        info_table.cell(idx, 0).paragraphs[0].runs[0].font.bold = True
        info_table.cell(idx, 0).paragraphs[0].runs[0].font.color.rgb = QU_MAROON
        info_table.cell(idx, 0).paragraphs[0].runs[0].font.size = Pt(11)
        
        # Style value cells
        info_table.cell(idx, 1).paragraphs[0].runs[0].font.size = Pt(11)
        info_table.cell(idx, 1).paragraphs[0].runs[0].font.color.rgb = RGBColor(51, 51, 51)
        
        # Background color
        set_cell_background(info_table.cell(idx, 0), 'F5F5F5')
        set_cell_background(info_table.cell(idx, 1), 'F5F5F5')
    
    # Set column widths
    info_table.columns[0].width = Inches(2)
    info_table.columns[1].width = Inches(8)
    
    # Add page numbering to footer
    for section in doc.sections:
        footer = section.footer
        footer_para = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
        footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        footer_para.text = "Page "
        
        # Add page number field
        run = footer_para.add_run()
        fldChar1 = OxmlElement('w:fldChar')
        fldChar1.set(qn('w:fldCharType'), 'begin')
        
        instrText = OxmlElement('w:instrText')
        instrText.set(qn('xml:space'), 'preserve')
        instrText.text = 'PAGE'
        
        fldChar2 = OxmlElement('w:fldChar')
        fldChar2.set(qn('w:fldCharType'), 'end')
        
        run._r.append(fldChar1)
        run._r.append(instrText)
        run._r.append(fldChar2)
        
        footer_para.add_run(' of ')
        
        # Add total pages field
        run2 = footer_para.add_run()
        fldChar3 = OxmlElement('w:fldChar')
        fldChar3.set(qn('w:fldCharType'), 'begin')
        
        instrText2 = OxmlElement('w:instrText')
        instrText2.set(qn('xml:space'), 'preserve')
        instrText2.text = 'NUMPAGES'
        
        fldChar4 = OxmlElement('w:fldChar')
        fldChar4.set(qn('w:fldCharType'), 'end')
        
        run2._r.append(fldChar3)
        run2._r.append(instrText2)
        run2._r.append(fldChar4)
        
        footer_para.runs[0].font.size = Pt(9)
        footer_para.runs[0].font.color.rgb = DARK_GRAY
    
    # Add page break - PLOs start on page 2
    doc.add_page_break()
    
    # Add PLOs section
    plo_heading = doc.add_heading('Program Learning Outcomes', level=2)
    plo_heading.runs[0].font.color.rgb = QU_MAROON
    plo_heading.runs[0].font.size = Pt(16)
    plo_heading.paragraph_format.space_after = Pt(10)
    
    for plo in data['plos']:
        p = doc.add_paragraph()
        p.add_run(f"{plo['code']}: ").font.bold = True
        p.add_run(f"{plo['code']}: ").font.color.rgb = QU_MAROON
        p.add_run(plo['description'])
        p.paragraph_format.space_after = Pt(8)
    
    doc.add_paragraph().paragraph_format.space_after = Pt(24)
    
    # Add mapping matrix heading
    matrix_heading = doc.add_heading('PLO-Competency Mapping Matrix', level=2)
    matrix_heading.runs[0].font.color.rgb = QU_MAROON
    matrix_heading.runs[0].font.size = Pt(16)
    matrix_heading.paragraph_format.space_after = Pt(12)
    
    # Build transposed matrix: PLOs as columns, competencies as rows with GA headers
    # Calculate total rows: 1 header + sum of (1 GA row + n competency rows for each GA)
    total_rows = 1  # Header row
    for ga in data['gas']:
        total_rows += 1 + len(ga['competencies'])  # GA row + competency rows
    
    total_cols = 2 + len(data['plos'])  # GA column + Competency column + PLO columns
    
    matrix_table = doc.add_table(rows=total_rows, cols=total_cols)
    matrix_table.style = 'Table Grid'
    
    # Header row
    matrix_table.cell(0, 0).text = 'Graduate Attributes'
    matrix_table.cell(0, 1).text = 'Supporting Competencies'
    
    for plo_idx, plo in enumerate(data['plos']):
        matrix_table.cell(0, 2 + plo_idx).text = plo['code']
    
    # Style header row
    for col_idx in range(total_cols):
        cell = matrix_table.cell(0, col_idx)
        set_cell_background(cell, '8B1538')  # Maroon
        cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)
        cell.paragraphs[0].runs[0].font.bold = True
        cell.paragraphs[0].runs[0].font.size = Pt(9)
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Data rows
    row_idx = 1
    for ga in data['gas']:
        # GA section row
        ga_cell = matrix_table.cell(row_idx, 0)
        ga_cell.text = f"{ga['code']}: {ga['name']}"
        
        # Merge first two columns for GA row
        ga_cell.merge(matrix_table.cell(row_idx, 1))
        
        # Style GA row
        set_cell_background(ga_cell, 'C8A882')  # Light gold
        ga_cell.paragraphs[0].runs[0].font.color.rgb = QU_MAROON
        ga_cell.paragraphs[0].runs[0].font.bold = True
        ga_cell.paragraphs[0].runs[0].font.size = Pt(9)
        
        # Empty cells for PLO columns in GA row
        for plo_idx in range(len(data['plos'])):
            plo_cell = matrix_table.cell(row_idx, 2 + plo_idx)
            set_cell_background(plo_cell, 'C8A882')
        
        row_idx += 1
        
        # Competency rows under this GA
        for comp in ga['competencies']:
            # Empty GA column
            matrix_table.cell(row_idx, 0).text = ''
            set_cell_background(matrix_table.cell(row_idx, 0), 'F5F5F5')
            
            # Competency name
            comp_cell = matrix_table.cell(row_idx, 1)
            comp_cell.text = f"{comp['code']} – {comp['name']}"
            set_cell_background(comp_cell, 'F5F5F5')
            comp_cell.paragraphs[0].runs[0].font.size = Pt(8)
            
            # Weights for each PLO
            for plo_idx, plo in enumerate(data['plos']):
                weight = plo['mappings'].get(comp['code'], '0.00')
                weight_cell = matrix_table.cell(row_idx, 2 + plo_idx)
                # Convert weight to string if it's a number
                weight_cell.text = str(weight) if isinstance(weight, (int, float)) else weight
                weight_cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                weight_cell.paragraphs[0].runs[0].font.size = Pt(8)
            
            row_idx += 1
    
    # Set column widths for mapping matrix - narrower PLO columns, wider competency titles
    # Apply to both table.columns AND individual cells for better enforcement
    matrix_table.columns[0].width = Inches(1.5)  # GA column
    matrix_table.columns[1].width = Inches(4.0)  # Competency column (increased)
    for plo_idx in range(len(data['plos'])):
        matrix_table.columns[2 + plo_idx].width = Inches(0.6)  # PLO columns (decreased)
    
    # Also set preferred width on each cell to force the width
    for row in matrix_table.rows:
        row.cells[0].width = Inches(1.5)
        row.cells[1].width = Inches(4.0)
        for plo_idx in range(len(data['plos'])):
            row.cells[2 + plo_idx].width = Inches(0.6)
    
    # Add summary
    doc.add_paragraph().paragraph_format.space_after = Pt(16)
    total_comps = sum(len(ga['competencies']) for ga in data['gas'])
    summary = doc.add_paragraph()
    summary.add_run('Total Mappings: ').font.bold = True
    summary.add_run(f"{data['total_mappings']} | ")
    summary.add_run('Total PLOs: ').font.bold = True
    summary.add_run(f"{len(data['plos'])} | ")
    summary.add_run('Total Competencies: ').font.bold = True
    summary.add_run(str(total_comps))
    
    doc.add_paragraph().paragraph_format.space_after = Pt(24)
    
    # Add justifications
    just_heading = doc.add_heading(f"Competency Justifications ({len(data['justifications'])})", level=2)
    just_heading.runs[0].font.color.rgb = QU_MAROON
    just_heading.runs[0].font.size = Pt(16)
    just_heading.paragraph_format.space_after = Pt(12)
    
    for just in data['justifications']:
        # Create a table for each justification
        just_table = doc.add_table(rows=1, cols=3)
        just_table.style = 'Table Grid'
        
        # Code cell
        code_cell = just_table.cell(0, 0)
        code_cell.text = just['competency_code']
        set_cell_background(code_cell, '8B1538')
        code_cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)
        code_cell.paragraphs[0].runs[0].font.bold = True
        code_cell.paragraphs[0].runs[0].font.size = Pt(9)
        code_cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Name cell
        name_cell = just_table.cell(0, 1)
        name_cell.text = just['competency_name']
        set_cell_background(name_cell, 'F5F5F5')
        name_cell.paragraphs[0].runs[0].font.color.rgb = QU_MAROON
        name_cell.paragraphs[0].runs[0].font.bold = True
        name_cell.paragraphs[0].runs[0].font.size = Pt(9)
        
        # Justification text cell
        text_cell = just_table.cell(0, 2)
        text_cell.text = just['text']
        text_cell.paragraphs[0].runs[0].font.size = Pt(9)
        text_cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(51, 51, 51)
        text_cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        
        # Set column widths - narrow code, reduce title, maximize justification
        # Apply to both table.columns AND individual cells for better enforcement
        just_table.columns[0].width = Inches(0.6)  # Reduced for code
        just_table.columns[1].width = Inches(1.8)  # Reduced for title
        just_table.columns[2].width = Inches(7.6)  # Maximized for justification
        
        # Also set preferred width on each cell to force the width
        just_table.rows[0].cells[0].width = Inches(0.6)
        just_table.rows[0].cells[1].width = Inches(1.8)
        just_table.rows[0].cells[2].width = Inches(7.6)
        
        # Add spacing after table
        doc.add_paragraph().paragraph_format.space_after = Pt(10)
    
    return doc

def main():
    if len(sys.argv) < 2:
        print(json.dumps({'error': 'No input data file provided'}))
        sys.exit(1)
    
    try:
        # Read JSON data from file (same as PDF export approach)
        data_file_path = sys.argv[1]
        with open(data_file_path, 'r', encoding='utf-8') as f:
            data = json.load(f)
        
        # Create document
        doc = create_mapping_document(data)
        
        # Save to output path
        import os
        output_path = data.get('output_path', os.path.join(os.path.dirname(os.path.dirname(__file__)), 'temp', 'mapping_output.docx'))
        doc.save(output_path)
        
        print(json.dumps({'success': True, 'output_path': output_path}))
    
    except Exception as e:
        import traceback
        print(json.dumps({'error': str(e), 'traceback': traceback.format_exc()}))
        sys.exit(1)

if __name__ == '__main__':
    main()
